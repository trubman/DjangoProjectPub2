from collections import namedtuple
from decimal import Decimal
from abc import ABC

from django.contrib import messages
from django.core.exceptions import ObjectDoesNotExist
from django.db.models.expressions import result
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Mm, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from sert.models import Sert, Kernel, Attachment, Melt, Guarantee, Conclusion, Signatories
from sert.models_data_bygost import ByGost


class SertIncarnation:
    def __init__(self, sert, kernel_data, attach_data, melt_data):
        # данные из model
        self.sert = sert
        self.kernel_data = kernel_data
        self.attach_data = attach_data
        self.melt_data = melt_data
        self.asymm_melts_rows = None
        self.sert_type = self.sert.sert_type
        # данные style
        ## основной внешний вид
        self.organization = None,
        self.paging = None,  # для создания страниц в docx SOLO/GROUP
        self.page_orientation = None,  # для работы setting_page DocxMaker PORTRAIT/LANDSCAPE
        self.head_image = None,  # для работы setting_header DocxMaker
        self.is_galv_date = None,
        self.melt_anatomy = None,  # для работы setting_body DocxMaker
        self.conclusion_type = None,  # для работы setting_conclusion DocxMaker
        self.sign_type = None,  # для работы setting_signs DocxMaker
        self.head_gost_str = None,  # для работы setting_header DocxMaker
        ## внешний вид таблиц
        self.tabs_names_list = []
        # данные tabs
        ## обязательные
        self.docx_tab = None
        self.main_tab = None
        ## факультативные
        self.is_symmetrical = True
        self.main_parts_tab = None
        self.parts_tab = None
        self.galv_parts_tab = None
        self.casts_tab = None
        self.chem_tab = None
        self.ctk_mech_tab = None
        self.ctk_chem_tab = None
        # данные docx
        self.docx = None
        ## для создания страниц в docx
        self.ctk_iter_index = None
        self.is_iter_index = False
        self.iter_index = 1
        self.sert_number_index = 1
        # для хранения результатов работы
        self.is_print_done = False
        self.errors_list = []  # (level, text,)


class MeltAnatomy:
    def __init__(self):
        self.metadata_list = self.generate_metadata_list()

    def get_metadata_by_material_id(self, id):
        return self.metadata_list[id]

    def generate_metadata_list(self):
        MeltMetaData = namedtuple('MeltMetaData',
                                  ['material_id', 'material_name', 'chem_list', 'mech_list', ])
        metadata_dict = {
            '32': MeltMetaData(
                material_id='32',
                material_name='20ГЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': ['tensile_strength', 'yield_strength', 'relative_extension',
                                     'relative_narrowing', 'impact_strength', 'hardness',],
                    'СЕРТ_ЦТК': ['tensile_strength', 'yield_strength', 'relative_extension',
                                 'relative_narrowing', 'impact_strength',],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': ['tensile_strength', 'yield_strength', 'relative_extension',
                                         'relative_narrowing', 'impact_strength_60KCV',],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': ['tensile_strength', 'yield_strength', 'relative_extension',
                                                  'relative_narrowing', 'impact_strength_60KCU'],
                }
            ),
            '32.1': MeltMetaData(
                material_id='32.1',
                material_name='35ГЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                 'impact_strength'],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '34': MeltMetaData(
                material_id='34',
                material_name='35ХМЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus',
                           'chromium', 'molybdaenum'],
                mech_list={
                    'АРМАТУРА_ХИМ': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                     'impact_strength', 'hardness'],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '35': MeltMetaData(
                material_id='35',
                material_name='20ХМЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus',
                           'chromium', 'molybdaenum'],
                mech_list={
                    'АРМАТУРА_ХИМ': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                     'impact_strength ', 'hardness'],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '35.1': MeltMetaData(
                material_id='35.1',
                material_name='20ГМЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus',
                           'chromium', 'molybdaenum', 'niccolum'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '16': MeltMetaData(
                material_id='16',
                material_name='12Х18Н9ТЛ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus',
                           'chromium', 'niccolum', 'titanium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                 'impact_strength'],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                 'impact_strength', 'mkk'],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '33': MeltMetaData(
                material_id='33',
                material_name='25Л',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': ['tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                 'impact_strength'],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '33.1': MeltMetaData(
                material_id='33.1',
                material_name='25Л',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [ 'tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                  'impact_strength'],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '33.2': MeltMetaData(
                material_id='33.2',
                material_name='35Л',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [ 'tensile_strength', 'yield_strength', 'relative_extension', 'relative_narrowing',
                                  'impact_strength'],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '22': MeltMetaData(
                material_id='22',
                material_name='СЧ20',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': ['tensile_strength',],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '02': MeltMetaData(
                material_id='02',
                material_name='06ХН28МДБ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'molybdaenum', 'niccolum', 'niobium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '05': MeltMetaData(
                material_id='05',
                material_name='06ХН28МДТ',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'molybdaenum', 'niccolum', 'titanium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '06': MeltMetaData(
                material_id='06',
                material_name='12Х18Н10Т',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'niccolum', 'titanium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '07': MeltMetaData(
                material_id='07',
                material_name='10Х17Н13М2Т',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'molybdaenum', 'niccolum', 'titanium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '13': MeltMetaData(
                material_id='13',
                material_name='12Х18Н10Б',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'niccolum', 'niobium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '14': MeltMetaData(
                material_id='14',
                material_name='10Х17Н13М2Б',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'molybdaenum', 'niccolum', 'niobium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '20': MeltMetaData(
                material_id='20',
                material_name='20Х13Л',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '44': MeltMetaData(
                material_id='44',
                material_name='АК7',
                chem_list=['manganum', 'silicium', 'cuprum', 'magnesium', 'ferrum'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '37': MeltMetaData(
                material_id='37',
                material_name='20Л',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
            '304L': MeltMetaData(
                material_id='304L',
                material_name='AiSi304L',
                chem_list=['carboneum', 'manganum', 'silicium', 'sulfur', 'phosphorus', 'chromium',
                           'niccolum'],
                mech_list={
                    'АРМАТУРА_ХИМ': [],
                    'СЕРТ_ЦТК': [],
                    'СЕРТ_ЦТК_НТ_ВЭЛВ': [],
                    'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': [],
                }
            ),
        }
        return metadata_dict


class SertStyleMaker:
    def __init__(self, sert_incarnation):
        self.si = sert_incarnation

        self.main_style_params = [
            'organization',
            'paging',  # для создания страниц в docx SOLO/GROUP
            'page_orientation',  # для работы setting_page DocxMaker PORTRAIT/LANDSCAPE
            'head_image',  # для работы setting_header DocxMaker
            'is_galv_date',
            'melt_anatomy',  # для работы setting_body DocxMaker
            'conclusion_type',  # для работы setting_conclusion DocxMaker
            'sign_type',  # для работы setting_signs DocxMaker
            'head_gost_str',  # для работы setting_header DocxMaker
        ]
        self.main_style_nt = namedtuple('main_style_nt', self.main_style_params)
        self.main_style = self.main_style_nt(
            organization='GG',
            paging='GROUP',
            page_orientation='PORTRAIT',
            head_image='static/serthead/header1.jpg',
            is_galv_date=False,
            melt_anatomy=None,
            conclusion_type='MAIN',
            sign_type='MAIN',
            head_gost_str='CTK_HEAD',
        )
        self.main_ctk_style = self.main_style_nt(
            organization='CTK',
            paging='SOLO',
            page_orientation='LANDSCAPE',
            head_image='static/serthead/litio1.jpg',
            is_galv_date=False,
            melt_anatomy=MeltAnatomy(),
            conclusion_type='CTK_CONC',
            sign_type='CTK',
            head_gost_str='CTK_HEAD',
        )
        self.main_styles_by_types = {
            'НАСОС': {'paging': 'SOLO', },
            'АРМАТУРА': {},
            'НАСОС_КУСОЧКИ': {'paging': 'SOLO', },
            'АРМАТУРА_КУСОЧКИ': {},
            'РЕМКОМПЛЕКТ': {},
            'ГАЛЬВАНИКА': {'is_galv_date': True, 'conclusion_type': 'GALV', 'sign_type': 'GALV', },
            'НАСОС_ХИМ': {'paging': 'SOLO', 'melt_anatomy': MeltAnatomy(), 'sign_type': 'CHEM1'},
            'АРМАТУРА_ХИМ': {'melt_anatomy': MeltAnatomy(), 'sign_type': 'CHEM1'},
            'СЕРТ_ЦТК': {},
            'СЕРТ_ЦТК_ГГ_НАСОС': {'paging': 'SOLO', 'melt_anatomy': MeltAnatomy(),
                                  'sign_type': 'CHEM_CTK', 'head_image': 'static/serthead/litio2.jpg', },
            'СЕРТ_ЦТК_НТ_ВЭЛВ': {},  # ???
            # 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': '',
        }
        self.data_styles_by_types = {
            'НАСОС': [],
            'АРМАТУРА': [],
            'НАСОС_КУСОЧКИ': ['main_parts_tab', ],
            'АРМАТУРА_КУСОЧКИ': ['main_parts_tab', ],
            'РЕМКОМПЛЕКТ': ['parts_tab', ],
            'ГАЛЬВАНИКА': ['galv_parts_tab', ],
            'НАСОС_ХИМ': ['casts_tab', 'chem_tab', ],
            'АРМАТУРА_ХИМ': ['casts_tab', 'chem_tab', ],
            'СЕРТ_ЦТК': ['ctk_mech_tab', 'ctk_chem_tab', ],
            'СЕРТ_ЦТК_ГГ_НАСОС': ['casts_tab', 'chem_tab', ],
            'СЕРТ_ЦТК_НТ_ВЭЛВ': ['ctk_mech_tab', 'ctk_chem_tab', ],  # ???
            # 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': '',
        }

        self.set_main_result_style()
        self.set_data_result_style()

    def set_main_result_style(self):

        def try_by_type(param):
            try:
                val = self.main_styles_by_types[self.si.sert_type][param]
            except KeyError:
                if self.si.sert_type not in ['СЕРТ_ЦТК', 'СЕРТ_ЦТК_НТ_ВЭЛВ', 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60']:
                    val = self.main_style.__getattribute__(param)
                else:
                    val = self.main_ctk_style.__getattribute__(param)
            return val

        for param in self.main_style_params:
            val = None
            if param not in ['conclusion_type', 'guarantee_type', 'sign_type']:
                val = try_by_type(param)
            else:
                val = self.si.sert.__getattribute__(param)
                if not val:
                    val = try_by_type(param)
                elif param == 'conclusion_type':
                    val = val.conclusion_type
            self.si.__setattr__(param, val)

    def set_data_result_style(self):
        self.si.tabs_names_list += self.data_styles_by_types[self.si.sert_type]

    def get_result(self):
        return self.si


class StaticTabsMaker:
    def __init__(self):
        self.ctk_result_nt = namedtuple('result_row',
                                        ['casts_head_row', 'casts_data_rows', 'chem_head_row', 'chem_data_rows',
                                         'asymm_docx_tab',])
        self.asymm_chem_result_nt = namedtuple('result_row',
                                               ['head_row', 'symm_data_rows', 'asymm_head_row', 'asymm_data_rows', ])
        self.asymm_result_nt = namedtuple('result_row',
                                          ['head_row', 'symm_data_rows', 'asymm_data_rows',])
        self.symm_result_nt = namedtuple('result_row', ['head_row', 'symm_data_rows',])

    @staticmethod
    def create_head_row(model_row, fields_names):
        head_row = model_row(**fields_names)
        return head_row

    @staticmethod
    def check_symmetrical(attachment_data):
        result = True
        cnt = 1
        for item in attachment_data:
            if item.n_index == cnt:
                result = False
                break
            cnt += 1
        return result

    def create_rows_for_main_parts_tab(self, kernel_data, attach_data, is_iter_index,
                         is_symmetrical, model_row, fields_list,):
        di = None
        li = None
        if not is_symmetrical:
            di = {}
            for i in range(1, kernel_data.quantity + 1):
                inner_data = []
                for ii in attach_data:
                    if ii.n_index == i:
                        inner_data.append(ii)
                inner_tab = self.create_row_for_main_parts_tab(is_iter_index, inner_data,
                                                 kernel_data, model_row,
                                                 fields_list,)
                di[i] = inner_tab
        else:
            li = []
            for i in self.create_row_for_main_parts_tab(is_iter_index, attach_data,
                                          kernel_data, model_row,
                                          fields_list,):
                li.append(i)
        return li, di

    @staticmethod
    def create_row_for_main_parts_tab(is_iter_index, attach_data, kernel_data, model_row, fields_list,):
        tab = []
        cnt = 1
        for row in attach_data:
            if not row.is_cast:
                if not is_iter_index:
                    quantity = row.quantity * kernel_data.quantity
                else:
                    quantity = row.quantity
                string = str(quantity)
                quantity = string.rstrip('0').rstrip('.') if '.' in string else string
                di = {}
                for n in fields_list:
                    if n in ['order_num']:
                        val = cnt
                    elif n in ['quantity']:
                        val = quantity
                    else:
                        val = row.__getattribute__(n)
                    di[n] = val
                tab.append(model_row(**di))
                cnt += 1
        return tab

    def create_rows_for_parts_tab(self, kernel_data, attach_data, model_row, fields_list,):
        result_tab = []
        Condition = namedtuple('Condition', [
            'is_first_selection',
            'data_parent_selection_condition',
            'data_child_selection_condition',
            'parent_comparison_condition',
            'child_comparison_condition',
        ])
        CONDITIONS = [
            Condition(
                is_first_selection=True,
                data_parent_selection_condition=None,
                data_child_selection_condition='is_one',
                parent_comparison_condition=None,
                child_comparison_condition=None,
            ),
            Condition(
                is_first_selection=False,
                data_parent_selection_condition='is_one',
                data_child_selection_condition='is_two',
                parent_comparison_condition='a_index',
                child_comparison_condition='b_index',
            ),
            Condition(
                is_first_selection=False,
                data_parent_selection_condition='is_two',
                data_child_selection_condition='is_three',
                parent_comparison_condition='a1_index',
                child_comparison_condition='b1_index',
            ),
            Condition(
                is_first_selection=False,
                data_parent_selection_condition='is_three',
                data_child_selection_condition='is_four',
                parent_comparison_condition='a2_index',
                child_comparison_condition='b2_index',
            ),
        ]
        for cond in CONDITIONS:
            data_parent = []
            if not cond.is_first_selection:
                for i in result_tab:
                    if i.__getattribute__(cond.data_parent_selection_condition):
                        data_parent.append(i)
            else:
                data_parent.append('plug')
            data_child = []
            for i in attach_data:
                if i.__getattribute__(cond.data_child_selection_condition):
                    data_child.append(i)
            tab = []
            for i in data_parent:
                cnt = 1
                for j in data_child:
                    if not cond.is_first_selection:
                        if (i.__getattribute__(cond.parent_comparison_condition) ==
                                j.__getattribute__(cond.child_comparison_condition)):
                            cnt_str = f'{i.order_num}.{cnt}'
                            quantity = j.quantity * Decimal(i.quantity)
                            tab = self.create_row_for_parts_tab(tab, j, cnt_str, quantity, model_row, fields_list,)
                            cnt += 1
                    else:
                        cnt_str = str(cnt)
                        quantity = j.quantity * kernel_data.quantity
                        tab = self.create_row_for_parts_tab(tab, j, cnt_str, quantity, model_row, fields_list,)
                        cnt += 1
            if tab:
                for i in tab:
                    result_tab.append(i)
        return result_tab

    @staticmethod
    def create_row_for_parts_tab(tab, row, cnt_str, quantity, model_row, fields_list,):
        di = {}
        string = str(quantity)
        quantity = string.rstrip('0').rstrip('.') if '.' in string else string
        for n in fields_list:
            if n in ['order_num']:
                val = cnt_str
            elif n in ['quantity']:
                val = quantity
            else:
                val = row.__getattribute__(n)
            di[n] = val
        tab.append(model_row(**di))
        return tab

    def create_rows_for_galv_parts_tab(self, kernel_data, attach_data, model_row, fields_list,):
        li = []
        for i in self.create_row_for_galv_parts_tab(attach_data, kernel_data, model_row, fields_list,):
            li.append(i)
        return li

    @staticmethod
    def create_row_for_galv_parts_tab(attach_data, kernel_data, model_row, fields_list,):
        tab = []
        cnt = 1
        for row in attach_data:
            quantity = row.quantity * kernel_data.quantity
            string = str(quantity)
            quantity = string.rstrip('0').rstrip('.') if '.' in string else string
            di = {}
            for n in fields_list:
                if n in ['order_num']:
                    val = cnt
                elif n in ['quantity']:
                    val = f'{quantity} {row.galvan_units}'
                else:
                    val = row.__getattribute__(n)
                di[n] = val
            tab.append(model_row(**di))
            cnt += 1
        return tab

    def create_rows_for_casts_tab(self, kernel_data, attach_data, is_iter_index,
                         is_symmetrical, model_row, fields_list, melt_anatomy,):
        di = None
        li = None
        dme = None
        if not is_symmetrical:
            di = {}
            dme = {}
            for i in range(1, kernel_data.quantity + 1):
                inner_data = []
                for ii in attach_data:
                    if ii.n_index == i:
                        inner_data.append(ii)
                inner_tab, inner_tab_melts = self.create_row_for_casts_tab(inner_data, is_iter_index, attach_data,
                                                 kernel_data, model_row,
                                                 fields_list, melt_anatomy,)
                di[i] = inner_tab
                dme[i] = inner_tab_melts
        else:
            li = []
            inner_data = None
            tab, tab_melts = self.create_row_for_casts_tab(inner_data, is_iter_index, attach_data,
                                          kernel_data, model_row,
                                          fields_list, melt_anatomy,)
            for i in tab:
                li.append(i)
        return li, di, dme

    @staticmethod
    def create_row_for_casts_tab(inner_data, is_iter_index, attach_data, kernel_data, model_row, fields_list, melt_anatomy,):
        tab = []
        tab_melts = []
        cnt = 1
        casts = []
        melts = []
        if inner_data:
            for row in inner_data:
                if not row.is_cast:
                    casts.append(row)
                else:
                    melts.append(row)
        else:
            for row in attach_data:
                if not row.is_cast:
                    casts.append(row)
                else:
                    melts.append(row)
        if inner_data and not melts:
            for row in attach_data:
                if row.is_cast:
                    melts.append(row)
        for row in casts:
            melts_for_row = []
            melts_for_row_str = []
            material_id_list = []
            for m in melts:
                if row.a_index == m.b_index:
                    melts_for_row.append(m)
                    if m.material_id not in material_id_list:
                        material_id_list.append(m.material_id)
                    melt_str = f'{m.melt_number}-{m.material_id}'
                    melts_for_row_str.append(melt_str)
            melts_one_str = ', '.join(melts_for_row_str)
            material_name = melt_anatomy.get_metadata_by_material_id(material_id_list[0]).material_name
            if not is_iter_index:
                quantity = row.quantity * kernel_data.quantity
            else:
                quantity = row.quantity
            string = str(quantity)
            quantity = string.rstrip('0').rstrip('.') if '.' in string else string
            di = {}
            for n in fields_list:
                if n in ['denomination']:
                    val = f'{row.__getattribute__(n)} (отливка)'
                elif n in ['material_name']:
                    val = material_name
                elif n in ['quantity']:
                    val = quantity
                elif n in ['melt']:
                    val = melts_one_str
                else:
                    val = row.__getattribute__(n)
                di[n] = val
            tab_melts += melts_for_row
            tab.append(model_row(**di))
            cnt += 1
        return tab, tab_melts

    @staticmethod
    def create_row_for_chem_tab(melt_data, model_row, fields_list,):
        tab = []
        cnt = 1
        for row in melt_data:
            di = {}
            for n in fields_list:
                if n in ['melt']:
                    val = f'{row.melt_number}-{row.material_id}'
                else:
                    val = row.__getattribute__(n)
                di[n] = val
            tab.append(model_row(**di))
            cnt += 1
        return tab

    @staticmethod
    def create_row_for_ctk_mech_tab(cast, melt_item, order_num_index, material_name, model_row, fields_list, mech_row,):
        di = {}
        for n in fields_list:
            if n in ['order_num']:
                val = order_num_index
            elif n in ['designation']:
                val = cast.designation
            elif n in ['quantity']:
                string = str(melt_item.melt_avatar.quantity)
                val = string.rstrip('0').rstrip('.') if '.' in string else string
            elif n in ['melt_number', 'melt']:
                val = f'{melt_item.melt_itself.melt_number}-{melt_item.melt_itself.material_id}'
            elif n in ['material_name']:
                val = material_name
            elif (mech_row is not None) and n in [
                'tensile_strength',  # Предел прочности
                'yield_strength',  # Предел текучести
                'relative_extension',  # Относительное удлинение
                'relative_narrowing',  # Относительное сужение
                'impact_strength',  # Ударная вязкость
                'impact_strength_60KCU',  # Ударная вязкость КСU-60
                'impact_strength_60KCV',  # Ударная вязкость КСV-60
                'hardness',  # Твердость
                'mkk',  # МКК
            ]:
                val = mech_row.__getattribute__(n)
            else:
                val = melt_item.melt_itself.__getattribute__(n)
                if not isinstance(val, type(Decimal(1))):
                    val = val
                else:
                    string = str(val)
                    val = string.rstrip('0').rstrip('.') if '.' in string else string
            di[n] = val
        result = model_row(**di)
        return result


class SertTabsMaker(StaticTabsMaker):
    def __init__(self, sert_incarnation):
        super().__init__()
        self.si = sert_incarnation
        self.create_tabs()

    def create_tabs(self):
        NAMES = {
            'main_parts_tab': 'create_main_parts_tab',
            'parts_tab': 'create_parts_tab',
            'galv_parts_tab': 'create_galv_parts_tab',
            'casts_tab': 'create_casts_tab',
            'chem_tab': 'create_chem_tab',
            'ctk_mech_tab': 'create_ctk_mech_tab',
            'ctk_chem_tab': 'create_ctk_chem_tab',
        }
        self.create_docx_tab()
        self.create_main_tab()
        for tab_name in self.si.tabs_names_list:
            self.__getattribute__(NAMES[tab_name]).__call__()

    # создать данные внешнего вида docx документа
    def create_docx_tab(self):
        field_list = [
            # для создания страниц в docx
            'is_iter_index',
            'iter_index',
            'sert_number_index',
            # для работы setting_page DocxMaker
            'page_width',
            'page_height',
            'page_orientation_flag',
            # для работы setting_header DocxMaker
            'head_image',
            'sert_number',
            'date',
            'head_gost_str',
            # для работы setting_conclusion DocxMaker
            'guarantee_text',
            'is_drag_met',
            'conclusion_text',
            # для работы setting_signs DocxMaker
            'signatories_list',
        ]
        row_model = namedtuple('docx_tab', field_list)
        di = {}

        # для создания страниц в docx
        if self.si.paging == 'SOLO':
            if self.si.sert_type not in ['СЕРТ_ЦТК', 'СЕРТ_ЦТК_НТ_ВЭЛВ', 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60',]:
                di['is_iter_index'] = True
                di['iter_index'] = self.si.kernel_data.quantity
            else:
                di['is_iter_index'] = True
                attach_rows = []
                for a in self.si.attach_data:
                    if (not a.melt_number and not a.material_id and
                        not a.melt_year and not a.melt_passport):
                        attach_rows.append(a)
                num = len(attach_rows)
                di['iter_index'] = num
        else:
            di['is_iter_index'] = False
            di['iter_index'] = 1
        di['sert_number_index'] = 1

        # для работы setting_page DocxMaker
        if self.si.page_orientation == 'PORTRAIT':
            di['page_width'] = Mm(210)
            di['page_height'] = Mm(297)
            di['page_orientation_flag'] = WD_ORIENT.PORTRAIT
        elif self.si.page_orientation == 'LANDSCAPE':
            di['page_width'] = Mm(297)
            di['page_height'] = Mm(210)
            di['page_orientation_flag'] = WD_ORIENT.LANDSCAPE

        # для работы setting_header DocxMaker
        di['head_image'] = self.si.head_image
        if not self.si.is_galv_date:
            date = self.si.sert.date
        else:
            date = self.si.sert.galvan_date
        sert_str_number = f'{self.si.sert.number_unique.number}'.zfill(4)
        dy = date.strftime('%y')
        sert_format_date = f'{date.day:02}/{date.month:02}/{dy}'
        sert_number = f'{sert_format_date}-{sert_str_number}'
        di['sert_number'] = sert_number
        di['date'] = date
        try:
            conclusion = Conclusion.objects.get(conclusion_type=self.si.head_gost_str)
            di['head_gost_str'] = conclusion.conclusion_text
        except ObjectDoesNotExist:
            error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                          f'ссылку на не предусмотренный ГОСТ для шапки сертификата - {self.si.head_gost_str}. '
                          f'Установлен тип "CTK_HEAD"')
            error = (messages.ERROR, error_text)
            self.si.errors_list.append(error)
            default_val = Conclusion.objects.get(conclusion_type='CTK_HEAD')
            di['head_gost_str'] = default_val.conclusion_text

        # для работы setting_conclusion DocxMaker
        if not self.si.sert.guarantee_type:
            di['guarantee_text'] = None
        else:
            guarantees = Guarantee.objects.filter(guarantee_type=self.si.sert.guarantee_type)
            if not guarantees:
                error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                              f'ссылку на не предусмотренный тип гарантии - {self.si.sert.guarantee_type}.')
                error = (messages.ERROR, error_text)
                self.si.errors_list.append(error)
                di['guarantee_text'] = None
            else:
                texts = []
                for t in guarantees:
                    texts.append(t.guarantee_text)
                di['guarantee_text'] = ', '.join(texts)
        di['is_drag_met'] = self.si.sert.is_drag_met
        if not self.si.conclusion_type:
            error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                          f'не содержит указания типа заключения. '
                          f'Установлен тип "MAIN"')
            error = (messages.ERROR, error_text)
            self.si.errors_list.append(error)
            default_val = Conclusion.objects.get(conclusion_type='MAIN')
            di['conclusion_text'] = default_val.conclusion_text
        else:
            try:
                conclusion = Conclusion.objects.get(conclusion_type=self.si.conclusion_type)
                di['conclusion_text'] = conclusion.conclusion_text
            except ObjectDoesNotExist:
                error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                              f'ссылку на не предусмотренный тип заключения - {self.si.conclusion_type}. '
                              f'Установлен тип "MAIN"')
                error = (messages.ERROR, error_text)
                self.si.errors_list.append(error)
                default_val = Conclusion.objects.get(conclusion_type='MAIN')
                di['conclusion_text'] = default_val.conclusion_text

        # для работы setting_signs DocxMaker
        if not self.si.sign_type:
            error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                          f'не содержит указания типа подписантов. '
                          f'Установлен тип "MAIN"')
            error = (messages.ERROR, error_text)
            self.si.errors_list.append(error)
            di['signatories_list'] = Signatories.objects.filter(sign_type='MAIN')
        else:
            signs = Signatories.objects.filter(sign_type=self.si.sign_type)
            if not signs:
                di['signatories_list'] = Signatories.objects.filter(sign_type='MAIN')
                error_text = (f'DataTabs.create_docx_tab: сертификат {self.si.sert.number_unique.id} '
                              f'ссылку на не предусмотренный тип подписантов - {self.si.sign_type}. '
                              f'Установлен тип "MAIN"')
                error = (messages.ERROR, error_text)
                self.si.errors_list.append(error)
            else:
                di['signatories_list'] = signs

        self.si.docx_tab = row_model(**di)
        self.upload_iter_index()

    def upload_iter_index(self):
        self.si.is_iter_index = self.si.docx_tab.is_iter_index
        self.si.iter_index = self.si.docx_tab.iter_index

    # создать таблицы с данными для DocxMaker из Models
    def create_main_tab(self):
        NT = namedtuple('main_tab', ['head_row', 'data_row',])
        # модель данных
        fields_list = [
            'denomination',
            'designation',
            'quantity',
            'number_spg',
        ]
        model_row = namedtuple('main_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = {
            'denomination': 'Наименование',
            'designation': 'Обозначение',
            'quantity': 'Кол-во\n(шт.)',
            'number_spg': 'Зав.№',
        }
        # строки с data
        head_row = model_row(**FIELDNAMES)
        data = self.si.kernel_data
        if not self.si.is_iter_index:
            quantity = data.quantity
        else:
            quantity = 1
        data_row = model_row(
            denomination=data.denomination,
            designation=data.designation,
            quantity=quantity,
            number_spg=data.number_spg,
        )
        self.si.main_tab = NT(head_row=head_row, data_row=data_row)

    def create_main_parts_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows', 'asymm_data_rows'
        # модель данных
        fields_list = [
            'order_num',
            'denomination',
            'designation',
            'quantity',
        ]
        model_row = namedtuple('main_parts_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = {
            'order_num': '№п/п',
            'denomination': 'Наименование',
            'designation': 'Обозначение',
            'quantity': 'Кол-во\n(шт.)',
        }
        # строки с data
        head_row = self.create_head_row(model_row, FIELDNAMES)
        kernel_data = self.si.kernel_data
        attach_data = self.si.attach_data
        is_iter_index = self.si.is_iter_index
        is_symmetrical = self.check_symmetrical(attach_data)
        self.si.is_symmetrical = is_symmetrical
        symm_data_rows, asymm_data_rows = self.create_rows_for_main_parts_tab(kernel_data, attach_data,
                                                                is_iter_index, is_symmetrical,
                                                                model_row, fields_list,)
        self.si.main_parts_tab = self.asymm_result_nt(
            head_row=head_row,
            symm_data_rows=symm_data_rows,
            asymm_data_rows=asymm_data_rows,
        )

    def create_parts_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows', 'asymm_data_rows'
        # модель данных
        fields_list = [
            'order_num',
            'denomination',
            'designation',
            'quantity',
            'a_index',
            'b_index',
            'a1_index',
            'b1_index',
            'a2_index',
            'b2_index',
            'is_one',
            'is_two',
            'is_three',
            'is_four',
        ]
        model_row = namedtuple('main_parts_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = {
            'order_num': '№п/п',
            'denomination': 'Наименование',
            'designation': 'Обозначение',
            'quantity': 'Кол-во\n(шт.)',
            'a_index': '',
            'b_index': '',
            'a1_index': '',
            'b1_index': '',
            'a2_index': '',
            'b2_index': '',
            'is_one': '',
            'is_two': '',
            'is_three': '',
            'is_four': '',
        }
        # строки с data
        head_row = self.create_head_row(model_row, FIELDNAMES)
        kernel_data = self.si.kernel_data
        attach_data = self.si.attach_data
        symm_data_rows = self.create_rows_for_parts_tab(kernel_data, attach_data, model_row, fields_list,)
        self.si.parts_tab = self.symm_result_nt(
            head_row=head_row,
            symm_data_rows=symm_data_rows,
        )

    def create_galv_parts_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows'
        # модель данных
        fields_list = [
            'order_num',
            'designation',
            'galvan_material',
            'quantity',
        ]
        model_row = namedtuple('main_parts_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = {
            'order_num': '№п/п',
            'designation': 'Наименование',
            'galvan_material': 'Покрытие',
            'quantity': 'Кол-во',
        }
        # строки с data
        head_row = self.create_head_row(model_row, FIELDNAMES)
        kernel_data = self.si.kernel_data
        attach_data = self.si.attach_data
        # is_iter_index = self.si.is_iter_index
        is_symmetrical = True
        self.si.is_symmetrical = is_symmetrical
        symm_data_rows = self.create_rows_for_galv_parts_tab(kernel_data, attach_data, model_row, fields_list,)
        self.si.galv_parts_tab = self.symm_result_nt(
            head_row=head_row,
            symm_data_rows=symm_data_rows,
        )

    def create_casts_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows', 'asymm_data_rows'
        # модель данных
        fields_list = [
            'denomination',
            'designation',
            'quantity',
            'material_name',
            'melt',
        ]
        model_row = namedtuple('main_parts_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = {
            'denomination': 'Наименование',
            'designation': 'Обозначение',
            'quantity': 'Кол-во\n(шт.)',
            'material_name': 'Материал',
            'melt': 'Плавка',
        }
        # строки с data
        head_row = self.create_head_row(model_row, FIELDNAMES)
        kernel_data = self.si.kernel_data
        attach_data = self.si.attach_data
        is_iter_index = self.si.is_iter_index
        is_symmetrical = self.check_symmetrical(attach_data)
        self.si.is_symmetrical = is_symmetrical
        symm_data_rows, asymm_data_rows, asymm_melts_rows = self.create_rows_for_casts_tab(kernel_data, attach_data,
                                                                                           is_iter_index,
                                                                                           is_symmetrical,
                                                                                           model_row, fields_list,
                                                                                           self.si.melt_anatomy,)
        if asymm_melts_rows:
            asymm_melts = {}
            for i in range(1, len(asymm_melts_rows) + 1):
                asymm_melts_list = []
                for m in asymm_melts_rows[i]:
                    melt_id = f'{m.melt_number}-{m.material_id}-{m.melt_year}-{m.melt_passport}'
                    try:
                        melt = Melt.objects.get(melt_id=melt_id)
                        asymm_melts_list.append(melt)
                    except ObjectDoesNotExist:
                        error_text = (f'SertTabsMaker: {melt_id} не найдена в Melt')
                        error = (messages.ERROR, error_text)
                        self.si.errors_list.append(error)
                asymm_melts[i] = asymm_melts_list
            self.si.asymm_melts_rows = asymm_melts
        self.si.casts_tab = self.asymm_result_nt(
            head_row=head_row,
            symm_data_rows=symm_data_rows,
            asymm_data_rows=asymm_data_rows,
        )

    def create_chem_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows', 'asymm_data_rows'
        head_row = None
        symm_data_rows = None
        asymm_head_row = None
        asymm_data_rows = None
        if (not self.si.is_symmetrical) and self.si.asymm_melts_rows:
            asymm_head_row = {}
            asymm_data_rows = {}
            melt_data = self.si.asymm_melts_rows
            for i in range(1, len(melt_data) + 1):
                melt_data_item = melt_data[i]
                head_row, data_rows = self.create_chem_tab_params(melt_data_item)
                asymm_head_row[i] = head_row
                asymm_data_rows[i] = data_rows
        else:
            melt_data = self.si.melt_data
            head_row, symm_data_rows = self.create_chem_tab_params(melt_data)
        self.si.chem_tab = self.asymm_chem_result_nt(
            head_row=head_row,
            symm_data_rows=symm_data_rows,
            asymm_head_row=asymm_head_row,
            asymm_data_rows=asymm_data_rows,
        )

    def get_fields_list_for_chem_tab(self, material_id_list, is_mech=False):
        field_names = []
        class StandartElem:
            def __init__(self, name):
                self.name = name
                self.count = 0
        field_names_standard = [
            StandartElem('carboneum'),
            StandartElem('manganum'),
            StandartElem('silicium'),
            StandartElem('sulfur'),
            StandartElem('phosphorus'),
            StandartElem('chromium'),
            StandartElem('molybdaenum'),
            StandartElem('niccolum'),
            StandartElem('niobium'),
            StandartElem('titanium'),
            StandartElem('cuprum'),
            StandartElem('magnesium'),
            StandartElem('ferrum'),
            StandartElem('mkk'),
            StandartElem('tensile_strength'),
            StandartElem('yield_strength'),
            StandartElem('relative_extension'),
            StandartElem('relative_narrowing'),
            StandartElem('impact_strength'),
            StandartElem('impact_strength_60KCU'),
            StandartElem('impact_strength_60KCV'),
            StandartElem('hardness'),
        ]
        materials_chem_lists = []
        for material_id in material_id_list:
            if not is_mech:
                materials_chem_lists.append(
                    self.si.melt_anatomy.get_metadata_by_material_id(material_id).chem_list
                )
            else:
                materials_chem_lists.append(
                    self.si.melt_anatomy.get_metadata_by_material_id(material_id).mech_list[self.si.sert_type]
                )
        for list in materials_chem_lists:
            for elem_name in list:
                for elem in field_names_standard:
                    if elem.name == elem_name:
                        elem.count += 1
        for elem in field_names_standard:
            if elem.count > 0:
                field_names.append(elem.name)
        return field_names

    def get_fieldnames_for_chem_tab(self, fields_list, is_mech=False):
        CHEMFIELDNAMES = {
            'melt': 'Плавка',
            'melt_passport': 'Паспорт',
            'carboneum': 'C',  # (углерод)
            'manganum': 'Mn',  # (марганец)
            'silicium': 'Si',  # (кремний)
            'sulfur': 'S',  # (сера)
            'phosphorus': 'P',  # (фосфор)
            'chromium': 'Cr',  # (хром)
            'molybdaenum': 'Mo',  # (молибден)
            'niccolum': 'Ni',  # (никель)
            'niobium': 'Nb',  # (ниобий)
            'titanium': 'Ti',  # (титан)
            'cuprum': 'Cu',  # (медь)
            'magnesium': 'Mg',  # (магний)
            'ferrum': 'Fe',  # (железо)
        }
        MECHFIELDNAMES = {
            'order_num': '№ п/п',
            'designation': 'Наименование изделия',
            'quantity': 'Кол-во',
            'melt_number': '№ плавки',
            'material_name': 'Марка стали',

            'mkk': 'МКК',  # МКК
            'tensile_strength': 'Предел прочности\nσв, МПа',  # Предел прочности
            'yield_strength': 'Предел текучести\nσт, МПа',  # Предел текучести
            'relative_extension': 'Относительное удлинение,\n%',  # Относительное удлинение
            'relative_narrowing': 'Относительное сужение\nψ, %',  # Относительное сужение
            'impact_strength': 'Ударная вязкость КСU,\nкДж/м²',  # Ударная вязкость
            'impact_strength_60KCU': 'Ударная вязкость\nКСU-60,\nкДж/м²',  # Ударная вязкость КСU-60
            'impact_strength_60KCV': 'Ударная вязкость\nКСV-60,\nкгс·м/см²',  # Ударная вязкость КСV-60
            'hardness': 'Твердость,\nHB',  # Твердость
        }
        FIELDNAMES = {}
        for name in fields_list:
            if not is_mech:
                FIELDNAMES[name] = CHEMFIELDNAMES[name]
            else:
                FIELDNAMES[name] = MECHFIELDNAMES[name]
        return FIELDNAMES

    def create_chem_tab_params(self, melt_data):
        material_id_list = []
        for m in melt_data:
            if m.material_id not in material_id_list:
                material_id_list.append(m.material_id)
        # модель данных
        pre_fields_list = ['melt', 'melt_passport',]
        post_fields_list = self.get_fields_list_for_chem_tab(material_id_list)
        fields_list = pre_fields_list + post_fields_list
        model_row = namedtuple('main_parts_tab_nt', fields_list)
        # технические характеристики строк
        FIELDNAMES = self.get_fieldnames_for_chem_tab(fields_list)
        # строки с data
        head_row = self.create_head_row(model_row, FIELDNAMES)
        melt_rows = self.create_row_for_chem_tab(melt_data, model_row, fields_list,)
        return head_row, melt_rows

    def create_ctk_mech_tab(self):
        # self.result_row: 'head_row', 'symm_data_rows', 'asymm_data_rows'
        casts_head_row = {}
        casts_data_rows = {}
        chem_head_row = {}
        chem_data_rows = {}
        asymm_docx_tab_nt = namedtuple('asymm_docx_tab', ['head_gost_str', 'conclusion_type',])
        asymm_docx_tab = {}
        attach_data = self.si.attach_data
        melt_data = self.si.melt_data
        cast_rows = []
        melt_rows = []
        for a in attach_data:
            if not a.is_cast:
                cast_rows.append(a)
            else:
                melt_rows.append(a)
        self.si.ctk_iter_index = len(cast_rows)
        cnt_index = 1
        for cast in cast_rows:
            m_nt = namedtuple('m_nt', ['melt_avatar', 'melt_itself',])
            melts_for_cast = []
            material_id_list = []
            for melt in melt_rows:
                if cast.a_index == melt.b_index:
                    melt_id = f'{melt.melt_number}-{melt.material_id}-{melt.melt_year}-{melt.melt_passport}'
                    for m in melt_data:
                        if m.melt_id == melt_id:
                            melts_for_cast.append(m_nt(
                                melt_avatar=melt,
                                melt_itself=m,
                            ))
                    if melt.material_id not in material_id_list:
                        material_id_list.append(melt.material_id)
            # для casts
            c_pre_fields_list = ['order_num', 'designation', 'quantity', 'melt_number', 'material_name',]
            c_post_fields_list = self.get_fields_list_for_chem_tab(material_id_list, is_mech=True)
            c_fields_list = c_pre_fields_list + c_post_fields_list
            c_model_row = namedtuple('c_model_row', c_fields_list)
            C_FIELDNAMES = self.get_fieldnames_for_chem_tab(c_fields_list, is_mech=True)
            # для melts
            m_pre_fields_list = ['melt', ]
            m_post_fields_list = self.get_fields_list_for_chem_tab(material_id_list)
            m_fields_list = m_pre_fields_list + m_post_fields_list
            m_model_row = namedtuple('m_model_row', m_fields_list)
            M_FIELDNAMES = self.get_fieldnames_for_chem_tab(m_fields_list)

            c_head_row = self.create_head_row(c_model_row, C_FIELDNAMES)
            m_head_row = self.create_head_row(m_model_row, M_FIELDNAMES)

            c_data_rows = []
            m_data_rows = []
            order_num_index = 1
            for melt_item in melts_for_cast:
                mech_row = None
                material_id = melt_item.melt_itself.material_id
                if melt_item.melt_avatar.is_by_gost_material_number:
                    bg = ByGost()
                    mech_row = bg.get_melt_mech_characteristics(
                        sert_type=self.si.sert_type,
                        material_id=material_id,
                        number=melt_item.melt_itself.by_gost_number,
                    )
                material_name = self.si.melt_anatomy.get_metadata_by_material_id(material_id).material_name
                c_result = self.create_row_for_ctk_mech_tab(
                    cast, melt_item, order_num_index, material_name, c_model_row, c_fields_list, mech_row,)
                m_result = self.create_row_for_ctk_mech_tab(
                    cast, melt_item, order_num_index, material_name, m_model_row, m_fields_list, mech_row,)
                c_data_rows.append(c_result)
                m_data_rows.append(m_result)
                order_num_index += 1

            casts_head_row[cnt_index] = c_head_row
            casts_data_rows[cnt_index] = c_data_rows
            chem_head_row[cnt_index] = m_head_row
            chem_data_rows[cnt_index] = m_data_rows
            try:
                m = material_id_list[0]
            except KeyError:
                m = None
            head_gost_str = None
            conclusion_type = None

            def get_conc_text(search_param):
                list_co = Conclusion.objects.filter(conclusion_type=search_param)
                if list_co:
                    CO = list_co[0]
                    result = CO.conclusion_text
                else:
                    result = f'Нет такого Conclusion: {search_param}'
                return result

            if self.si.sert_type == 'СЕРТ_ЦТК_НТ_ВЭЛВ' and (m in ['32', '16']):
                if m == '32':
                    head_gost_str = get_conc_text('CTK_HEAD1')
                    conclusion_type = get_conc_text('CTK_CONC1')
                elif m == '16':
                    head_gost_str = get_conc_text('CTK_HEAD2')
                    conclusion_type = get_conc_text('CTK_CONC2')
            asymm_docx_tab[cnt_index] = asymm_docx_tab_nt(head_gost_str=head_gost_str, conclusion_type=conclusion_type)
            cnt_index += 1
        self.si.ctk_mech_tab = self.ctk_result_nt(
            casts_head_row=casts_head_row,
            casts_data_rows=casts_data_rows,
            chem_head_row=chem_head_row,
            chem_data_rows=chem_data_rows,
            asymm_docx_tab=asymm_docx_tab,
        )

    def create_ctk_chem_tab(self):
        pass

    def get_result(self):
        return self.si


class ModelDataLoader:

    @staticmethod
    def get_kernel_data(sert):
        data = sert.number_spg
        return data

    @staticmethod
    def get_attachment_data(sert):
        number_spg = sert.number_spg
        sert_type = sert.sert_type
        number_sert = sert.number_unique
        test = Attachment.objects.filter(number_unique=number_sert)
        if not test:
            data = (Attachment.objects.filter(number_spg=number_spg)
                    .filter(sert_type=sert_type))
        else:
            data = (Attachment.objects.filter(number_unique=number_sert)
                    .filter(number_spg=number_spg)
                    .filter(sert_type=sert_type))
        return data

    @staticmethod
    def get_melt_data(attachment_data_item):
        melt_number = attachment_data_item.melt_number
        material_id = attachment_data_item.material_id
        melt_year = attachment_data_item.melt_year
        melt_passport = attachment_data_item.melt_passport
        melt_id = f'{melt_number}-{material_id}-{melt_year}-{melt_passport}'
        data = Melt.objects.filter(melt_id=melt_id)
        return data


class DocxMakerStatic(ABC):

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def add_repeat_table_header(table):
        from docx.oxml import OxmlElement
        tbl_header = OxmlElement('w:tblHeader')  # create new oxml element flag which indicates that row is header row
        first_row_props = table.rows[0]._element.get_or_add_trPr()  # get if exists or create new table row properties el
        first_row_props.append(tbl_header)  # now first row is the header row
        return table

    # to change background colour of the column cells
    @staticmethod
    def change_columns_background(row_count, *columns, color='grey1'):
        COLORS = {
            'grey1': r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w')),
            'grey2': r'<w:shd {} w:fill="C0C0C0"/>'.format(nsdecls('w')),
            'grey3': r'<w:shd {} w:fill="A0A0A0"/>'.format(nsdecls('w')),
        }
        for i in range(0, row_count + 1):
            shading_elm_1 = parse_xml(COLORS[color])
            columns[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)

    # ???
    # расширять/сужать строки в зависимости от длинны таблицы
    @staticmethod
    def tune_up_table_rows_heights(table):
        len_tab = len(table.rows) - 1
        if ((21 < len_tab) and (len_tab < 33)) or ((68 < len_tab) and (len_tab < 82)):
            count = len(table.rows)
            while count >= 0:
                table.rows[count - 1].height = Mm(10)
                count -= 1

    @staticmethod
    def data_tab_set_head_row(table, head_row_data):
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row_data)):
            hdr_cells[i].text = head_row_data[i]

    @staticmethod
    def style_tab_set_widths_for_cols(table, widths=None):
        if not widths:
            ## всего 190 Mm
            widths = [Mm(10), Mm(80), Mm(40), Mm(50), Mm(10)]
        # (LibreOffice) задать размер колонок таблицы через настройку ячеек
        for i in range(len(table.rows)):
            cells = table.rows[i].cells
            for j in range(len(table.columns)):
                cells[j].width = widths[j]
        # (MSOffice) через настройку столбцов
        for i in range(len(table.columns)):
            table.columns[i].width = widths[i]

    @staticmethod
    def style_tabrow_set_alignment_and_font_bold(row, font_bold=False, is_change_font_size=False, font_size=12):
        row_cells = row.cells
        count = len(row_cells)
        while count >= 0:
            row_cells[count - 1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for i in row_cells[count - 1].paragraphs:
                i.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if font_bold:
                    for ii in i.runs:
                        ii.font.bold = True
                if is_change_font_size:
                    for ii in i.runs:
                        ii.font.size = Pt(font_size)
            count -= 1


class DocxMaker(DocxMakerStatic):
    def __init__(self, sert_incarno, docx):
        self.si = sert_incarno
        self.docx = docx
        self.make_frankenstein()

    def get_docx(self):
        return self.docx

    # настроить docx добавлением информации
    def make_frankenstein(self):
        self.setting_page()

        if self.si.sert_type not in ['СЕРТ_ЦТК', 'СЕРТ_ЦТК_НТ_ВЭЛВ', 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60']:

            if self.si.sert_type in ['РЕМКОМПЛЕКТ']:
                self.setting_header_for_parts_tab()
            else:
                self.setting_header()

            if self.si.sert_type in ['НАСОС', 'АРМАТУРА']:
                self.setting_main_tab_gross()
            elif self.si.sert_type in ['НАСОС_КУСОЧКИ', 'АРМАТУРА_КУСОЧКИ', 'РЕМКОМПЛЕКТ',
                                       'НАСОС_ХИМ', 'АРМАТУРА_ХИМ', 'СЕРТ_ЦТК_ГГ_НАСОС',]:
                self.setting_main_tab_klein()

            if self.si.sert_type not in ['НАСОС', 'АРМАТУРА']:
                self.setting_body()
        else:
            self.setting_header_for_ctk()
            self.setting_body()

        self.setting_conclusion()
        self.setting_signs()

        self.set_errors_strings()

        if self.si.iter_index != self.si.sert_number_index:
            self.docx.add_page_break()

    def set_errors_strings(self):
        if self.si.errors_list:
            prf = self.docx.add_paragraph(' ')
            prf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for error in self.si.errors_list:
                e_prf = self.docx.add_paragraph(error.text)
                e_prf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for i in e_prf.runs:
                    i.font.size = Pt(20)
                    i.font.color.rgb = RGBColor.from_string('8B0000')
            prf = self.docx.add_paragraph(' ')
            prf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def setting_body(self):
        METHODS = {
            'main_parts_tab': 'setting_body_for_main_parts_tab',
            'parts_tab': 'setting_body_for_parts_tab',
            'galv_parts_tab': 'setting_body_for_galv_parts_tab',
            'casts_tab': 'setting_body_for_casts_tab',
            'chem_tab': 'setting_body_for_chem_tab',
            'ctk_mech_tab': 'setting_body_for_ctk_mech_tab',
            'ctk_chem_tab': 'setting_body_for_ctk_chem_tab',
        }
        for tab_name in self.si.tabs_names_list:
            self.__getattribute__(METHODS[tab_name]).__call__()

    def setting_page(self):
        # определи базовые настройки шрифта (имя и размер)
        style_all_document = self.docx.styles['Normal']
        style_all_document.font.name = 'Times New Roman'
        style_all_document.font.size = Pt(12)
        paragraph_format = self.docx.styles['Normal'].paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # интервал между абзацами
        paragraph_format.space_after = Pt(0)
        # определи поля документа
        section = self.docx.sections[0]
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        # задай настройки страницы - её размер и ориентацию
        self.docx.sections[0].page_width = self.si.docx_tab.page_width
        self.docx.sections[0].page_height = self.si.docx_tab.page_height
        self.docx.sections[0].orientation = self.si.docx_tab.page_orientation_flag

    def create_sert_number_str(self):
        if not self.si.is_iter_index:
            sert_number_str = f'{self.si.docx_tab.sert_number}'
        else:
            sert_number_str = f'{self.si.docx_tab.sert_number}-{self.si.sert_number_index}'
        return sert_number_str

    def add_post_paragraph(self):
        if self.si.sert_type not in ['ГАЛЬВАНИКА']:
            str_empty_date = ' '
            font_size = 6
        else:
            str_empty_date = self.si.docx_tab.date.strftime('%d.%m.%Y')
            font_size = 12
        post_paragraph = self.docx.add_paragraph(str_empty_date)
        post_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        post_paragraph.paragraph_format.line_spacing = Pt(font_size)
        for i in post_paragraph.runs:
            i.font.size = Pt(font_size)

    def setting_header(self):
        self.docx.add_picture(self.si.docx_tab.head_image, width=Mm(190))
        picture_paragraph = self.docx.paragraphs[-1]
        picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sert_number_str = self.create_sert_number_str()
        sert_number_paragraph = self.docx.add_paragraph(f'СЕРТИФИКАТ КАЧЕСТВА №{sert_number_str}')
        sert_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in sert_number_paragraph.runs:
            i.font.bold = True
            i.font.size = Pt(18)
        self.add_post_paragraph()

    def setting_header_for_parts_tab(self):
        post_paragraph = self.docx.add_paragraph(str(self.si.main_tab.data_row.number_spg))
        for i in post_paragraph.runs:
            i.font.size = Pt(12)
            i.font.color.rgb = RGBColor.from_string('C0C0C0')
        self.setting_header()
        post_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def get_data_str_for_ctk(self, data):
        MN = {
            1: 'января',
            2: 'февраля',
            3: 'марта',
            4: 'апреля',
            5: 'мая',
            6: 'июня',
            7: 'июля',
            8: 'августа',
            9: 'сентября',
            10: 'октября',
            11: 'ноября',
            12: 'декабря',
        }
        d = data.day
        m = data.month
        y = data.year
        data_str = f'"{d:02}" {MN[m]} {y} г.'
        return data_str

    def setting_header_for_ctk(self):
        prf = self.docx.add_paragraph(' ')
        prf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i in prf.runs:
            i.font.size = Pt(6)

        table = self.docx.add_table(rows=1, cols=4)
        # table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = [Mm(10), Mm(30), Mm(45), Mm(192)]
        self.style_tab_set_widths_for_cols(table, widths)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        cell_icon = hdr_cells[1].add_paragraph()
        cell_icon.add_run()
        cell_icon.runs[0].add_picture(self.si.head_image, width=Mm(25))
        self.delete_paragraph(hdr_cells[1].paragraphs[0])
        hdr_cells[2].text = 'Россия,\nВоронежская обл.,\nПанинский р-он,\nООО «ЦТК «Литьё»'
        hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i in hdr_cells[2].paragraphs:
            for ii in i.runs:
                ii.font.bold = True
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sv = self.si.docx_tab.sert_number.split('-')[1]
        hdr_cells[3].paragraphs[0].text = (
            f'Сертификат качества №{sv}-{self.si.sert_number_index}')
        for i in hdr_cells[3].paragraphs:
            for ii in i.runs:
                # ii.font.bold = True
                ii.font.size = Pt(18)
        if self.si.sert_type in ['СЕРТ_ЦТК_НТ_ВЭЛВ', 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60']:
            head_gost_str = self.si.ctk_mech_tab.asymm_docx_tab[self.si.sert_number_index].head_gost_str
            if head_gost_str is None:
                head_gost_str = self.si.docx_tab.head_gost_str
        else:
            head_gost_str = self.si.docx_tab.head_gost_str
        prf1 = hdr_cells[3].add_paragraph(head_gost_str)
        prf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prf1.runs[0].font.size = Pt(12)
        hdr_cells[3].add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        prf3 = hdr_cells[3].add_paragraph(self.get_data_str_for_ctk(self.si.docx_tab.date))
        prf3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prf3.runs[0].font.size = Pt(12)

        prf = self.docx.add_paragraph(self.si.main_tab.data_row.number_spg)
        prf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prf.paragraph_format.line_spacing = Pt(8)
        for i in prf.runs:
            i.font.size = Pt(8)
            i.font.color.rgb = RGBColor.from_string('C0C0C0')

    def setting_main_tab_gross(self):
        name_paragraph = self.docx.add_paragraph('Наименование продукции:')
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table = self.docx.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # заполни шапку талицы
        head_row_data = self.si.main_tab.head_row
        self.data_tab_set_head_row(table, head_row_data)
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)

        # создай, заполни и настрой единственную строку
        data_row = self.si.main_tab.data_row
        row_cells = table.rows[1].cells
        for i in range(len(data_row)):
            val = data_row[i]
            if val == data_row.number_spg:
                if self.si.is_iter_index:
                    row_cells[i].text = f'{val}-{self.si.sert_number_index:02}'
                else:
                    if data_row.quantity > 1:
                        row_cells[i].text = f'{val}-01÷{data_row.quantity:02}'
                    elif data_row.quantity == 1:
                        row_cells[i].text = f'{val}-01'
            elif val == data_row.quantity:
                if self.si.is_iter_index:
                    row_cells[i].text = str(1)
                else:
                    row_cells[i].text = str(val)
            else:
                row_cells[i].text = str(val)
        self.style_tabrow_set_alignment_and_font_bold(table.rows[1])

        # задать размер колонок таблицы
        ## всего 190 Mm
        widths = [Mm(40), Mm(90), Mm(20), Mm(40)]
        self.style_tab_set_widths_for_cols(table, widths)

    def generate_prod_name_for_name_table(self):
        root_prod_name = ''
        main_tab = self.si.main_tab.data_row
        if self.si.sert_type not in ['РЕМКОМПЛЕКТ']:
            if self.si.is_iter_index:
                root_prod_name = f'{main_tab.designation} зав.№{main_tab.number_spg}-{self.si.sert_number_index:02}'
            else:
                if main_tab.quantity > 1:
                    root_prod_name = f'{main_tab.designation} зав.№{main_tab.number_spg}-01÷{main_tab.quantity:02}'
                elif main_tab.quantity == 1:
                    root_prod_name = f'{main_tab.designation} зав.№{main_tab.number_spg}-01'
        else:
            root_prod_name = f'{main_tab.designation}'
        return root_prod_name

    def setting_main_tab_klein(self):
        root_prod_name = self.generate_prod_name_for_name_table()
        name_table = self.docx.add_table(rows=1, cols=2)
        # name_table.style = 'Table Grid'
        name_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # задать размер колонок таблицы
        for i in range(len(name_table.rows)):
            cells = name_table.rows[i].cells
            # всего 190 Mm
            # widths = [Mm(55), Mm(135), ]
            for j in range(len(name_table.columns)):
                # cells[j].width = widths[j]
                cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        widths = [Mm(55), Mm(135),]
        self.style_tab_set_widths_for_cols(name_table, widths)

        name_row_cells = name_table.rows[0].cells
        name_row_cells[0].text = 'Наименование продукции:'
        for i in name_row_cells[0].paragraphs:
            i.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_row_cells[1].text = f'{root_prod_name}'
        for i in name_row_cells[1].paragraphs:
            i.alignment = WD_ALIGN_PARAGRAPH.LEFT

        end_pr = self.docx.add_paragraph(' ')
        end_pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        end_pr.paragraph_format.line_spacing = Pt(1)
        for i in end_pr.runs:
            i.font.size = Pt(1)

    def setting_conclusion(self):
        if self.si.docx_tab.guarantee_text:
            guarantee_paragraph = self.docx.add_paragraph(self.si.docx_tab.guarantee_text)
            guarantee_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            guarantee_paragraph.paragraph_format.first_line_indent = Mm(10)
            for i in guarantee_paragraph.runs:
                i.font.size = Pt(12)
        if self.si.docx_tab.is_drag_met:
            drag_met_paragraph = self.docx.add_paragraph('Драгоценные металлы отсутствуют.')
            drag_met_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            drag_met_paragraph.paragraph_format.first_line_indent = Mm(10)
            for i in drag_met_paragraph.runs:
                i.font.bold = True
                i.font.size = Pt(12)
        if self.si.sert_type in ['СЕРТ_ЦТК_НТ_ВЭЛВ', 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60']:
            conclusion_text = self.si.ctk_mech_tab.asymm_docx_tab[self.si.sert_number_index].conclusion_type
            if conclusion_text is None:
                conclusion_text = self.si.docx_tab.conclusion_text
        else:
            conclusion_text = self.si.docx_tab.conclusion_text
        conclusion_paragraph = self.docx.add_paragraph(conclusion_text)
        conclusion_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        conclusion_paragraph.paragraph_format.first_line_indent = Mm(10)
        for i in conclusion_paragraph.runs:
            i.font.size = Pt(12)
        if self.si.organization == 'GG':
            empty_paragraph = self.docx.add_paragraph(' ')
            empty_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            empty_paragraph = self.docx.add_paragraph(' ')
            empty_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.si.organization == 'CTK':
            empty_paragraph = self.docx.add_paragraph(' ')
            empty_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            empty_paragraph.paragraph_format.line_spacing = Pt(8)

    def setting_signs(self):
        rows = len(self.si.docx_tab.signatories_list) + (len(self.si.docx_tab.signatories_list) - 1)
        table = self.docx.add_table(rows=rows, cols=5)
        # table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cnt = 0
        empty_cnt = 1
        for i in range(rows):
            sign = self.si.docx_tab.signatories_list[cnt]
            if not empty_cnt % 2 == 0:
                row_cells = table.rows[i].cells
                for i in row_cells:
                    i.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                row_cells[1].text = sign.sign_job_title
                for i in row_cells[1].paragraphs:
                    i.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row_cells[2].text = '_________________'
                row_cells[3].text = sign.sign_person
                cnt += 1
            empty_cnt += 1
        widths = [Mm(10), Mm(80), Mm(40), Mm(50), Mm(10)]
        if self.si.organization == 'CTK':
            widths = [Mm(53), Mm(80), Mm(40), Mm(50), Mm(54)]
        self.style_tab_set_widths_for_cols(table, widths)

    def setting_body_for_main_parts_tab(self):
        # взять данные для заполнения таблицы в зависимости от симметрии
        head_row = self.si.main_parts_tab.head_row
        if self.si.is_symmetrical:
            data_rows = self.si.main_parts_tab.symm_data_rows
        else:
            data_rows = self.si.main_parts_tab.asymm_data_rows[self.si.sert_number_index]
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)
        # заполнить тело таблицы
        for row in data_rows:
            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # настрой ширины колонок таблицы
        widths = [Mm(10), Mm(60), Mm(100), Mm(20), ]
        self.style_tab_set_widths_for_cols(table, widths)

    def sort_data_by_attach_index(self, data):
        SD = namedtuple(
            'SD',
            ['is_one', 'is_two', 'is_three', 'is_four',]
        )
        sorted_data = SD(is_one=[], is_two=[], is_three=[], is_four=[])
        for i in data:
            if i.is_one:
                sorted_data.is_one.append(i)
            elif i.is_two:
                sorted_data.is_two.append(i)
            elif i.is_three:
                sorted_data.is_three.append(i)
            elif i.is_four:
                sorted_data.is_four.append(i)
        return sorted_data

    def fill_tab_row(self, row, data_row, fontsize=12):

        def tune_up_fontsize_in_cell(cell, size):
            for i in cell.paragraphs:
                for ii in i.runs:
                    ii.font.size = Pt(size)

        row_cells = row.cells
        row_cells[0].text = str(data_row.order_num)
        tune_up_fontsize_in_cell(row_cells[0], fontsize)
        row_cells[1].text = str(data_row.denomination)
        row_cells[2].text = str(data_row.designation)
        row_cells[3].text = str(data_row.quantity)

    def get_is_need_inner_row(self, row, sorted_data, nesting_index):
        Nesting = namedtuple('Nesting', [
            'sorted_data_attr', 'a_index', 'b_index',
        ])
        NESTING = {
            1: Nesting(sorted_data_attr='is_two', a_index='a_index', b_index='b_index',),
            2: Nesting(sorted_data_attr='is_three', a_index='a1_index', b_index='b1_index', ),
            3: Nesting(sorted_data_attr='is_four', a_index='a2_index', b_index='b2_index', ),
        }
        is_need_inner_row = False
        for i in sorted_data.__getattribute__(NESTING[nesting_index].sorted_data_attr):
            if (row.__getattribute__(NESTING[nesting_index].a_index)
                    == i.__getattribute__(NESTING[nesting_index].b_index)):
                is_need_inner_row = True
                break
        return is_need_inner_row

    def add_row_for_nested_table_and_table_itself(self, table):

        def merge_four_cells_in_one(table_row):
            row_cells_for_merge = table_row.cells
            row_cells_for_merge[0].merge(row_cells_for_merge[1])
            row_cells_for_merge[0].merge(row_cells_for_merge[2])
            row_cells_for_merge[0].merge(row_cells_for_merge[3])
            row_cells_for_merge[0].text = 'в составе:'
            row_cells_for_merge[0].paragraphs[0].runs[0].font.size = Pt(8)
            return table_row

        current_table_row = table.add_row()
        current_table_row = merge_four_cells_in_one(current_table_row)
        inner_table = current_table_row.cells[0].add_table(rows=1, cols=4)
        inner_table.style = 'Table Grid'
        inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        current_table_row.cells[0].paragraphs[1].add_run(' ')
        current_table_row.cells[0].paragraphs[1].paragraph_format.line_spacing = Pt(6)
        return inner_table

    def tune_up_table_columns_widths_and_color(self, table, widths_index=0, color='grey1'):
        WIDTHS = {
            0: None,
            1: [Mm(18), Mm(55), Mm(95), Mm(18), ],
            2: [Mm(16), Mm(55), Mm(95), Mm(16), ],
            3: [Mm(14), Mm(55), Mm(95), Mm(14), ],
        }
        if table:
            # задать размер колонок таблицы
            self.style_tab_set_widths_for_cols(table, WIDTHS[widths_index])

            row_count = len(table.rows) - 1
            for i in range(4):
                self.change_columns_background(row_count, table.columns[i], color=color)

    def fill_inner_tab(self, data_row, inner_table, is_first_row, fontsize=12):
        if is_first_row:
            first_inner_row = inner_table.rows[0]
            self.fill_tab_row(first_inner_row, data_row, fontsize)
            self.style_tabrow_set_alignment_and_font_bold(first_inner_row)
        else:
            inner_row = inner_table.add_row()
            self.fill_tab_row(inner_row, data_row, fontsize)
            self.style_tabrow_set_alignment_and_font_bold(inner_row)
        is_first_row = False
        return is_first_row

    def setting_body_for_parts_tab(self):
        # взять данные для заполнения таблицы в зависимости от глубины вложенности
        head_row = self.si.parts_tab.head_row
        head_row = head_row[:len(head_row) - 10]
        data_rows = self.si.parts_tab.symm_data_rows
        sorted_data = self.sort_data_by_attach_index(data_rows)
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)
        # заполнить тело таблицы
        # построчно заполняет основную таблицу №1
        for main_row in sorted_data.is_one:
            # добавляет строку в основную таблицу и заполняет её содержимым
            current_table_row1 = table.add_row()
            self.fill_tab_row(current_table_row1, main_row)
            # определяет есть ли у данной строки вложенные подстроки
            is_need_inner_row = self.get_is_need_inner_row(main_row, sorted_data, 1)
            inner_table2 = None
            is_first_row2 = True
            # добавляет в основную таблицу пустую строку, сращивает её в одну и создает внутри вложенную таблицу
            if is_need_inner_row:
                inner_table2 = self.add_row_for_nested_table_and_table_itself(table)
            # построчно заполняет вложенную таблицу №2
            for row2 in sorted_data.is_two:
                # отбирает из сортированного списка те строки, которые имеют отношение к данной вложенной строке
                if main_row.a_index == row2.b_index:

                    is_first_row2 = self.fill_inner_tab(row2, inner_table2, is_first_row2)

                    is_need_inner_row2 = self.get_is_need_inner_row(row2, sorted_data, 2)
                    inner_table3 = None
                    is_first_row3 = True

                    if is_need_inner_row2:
                        inner_table3 = self.add_row_for_nested_table_and_table_itself(inner_table2)
                    # построчно заполняет вложенную таблицу №3
                    for row3 in sorted_data.is_three:
                        if row2.a1_index == row3.b1_index:
                            is_first_row3 = self.fill_inner_tab(row3, inner_table3, is_first_row3, fontsize=10)

                            is_need_inner_row3 = self.get_is_need_inner_row(row3, sorted_data, 3)
                            inner_table4 = None
                            is_first_row4 = True
                            if is_need_inner_row3:
                                inner_table4 = self.add_row_for_nested_table_and_table_itself(
                                    inner_table3)
                            # построчно заполняет вложенную таблицу №4
                            for row4 in sorted_data.is_four:
                                if row3.a2_index == row4.b2_index:
                                    is_first_row4 = self.fill_inner_tab(row4, inner_table4, is_first_row4, fontsize=10)

                            self.tune_up_table_columns_widths_and_color(inner_table4, widths_index=3, color='grey3')

                    self.tune_up_table_columns_widths_and_color(inner_table3, widths_index=2, color='grey2')
            # задает размер и цвет колонкам строки
            self.tune_up_table_columns_widths_and_color(inner_table2, widths_index=1)
            # задает выравнивание колонкам строки
            self.style_tabrow_set_alignment_and_font_bold(current_table_row1)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # задать размер колонок основной таблицы
        widths = [Mm(20), Mm(55), Mm(95), Mm(20),]
        self.style_tab_set_widths_for_cols(table, widths)
        # управляет высотой строк в зависимости от их количества
        # это сомнительный механизм для корректного заполнения страниц документа
        # чтобы одинокая подпись не слетала на новую страницу, а имела хоть какие-то
        # строки таблицы над собой
        self.tune_up_table_rows_heights(table)

    def setting_body_for_galv_parts_tab(self):
        name_paragraph = self.docx.add_paragraph('Наименование продукции:')
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # взять данные для заполнения таблицы в зависимости от симметрии
        head_row = self.si.galv_parts_tab.head_row
        data_rows = self.si.galv_parts_tab.symm_data_rows
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)
        # заполнить тело таблицы
        for row in data_rows:
            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # настрой ширины колонок таблицы
        widths = [Mm(20), Mm(110), Mm(30), Mm(30),]
        self.style_tab_set_widths_for_cols(table, widths)

    def setting_body_for_casts_tab(self):
        # взять данные для заполнения таблицы в зависимости от симметрии
        head_row = self.si.casts_tab.head_row
        if self.si.is_symmetrical:
            data_rows = self.si.casts_tab.symm_data_rows
        else:
            data_rows = self.si.casts_tab.asymm_data_rows[self.si.sert_number_index]
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)
        # заполнить тело таблицы
        for row in data_rows:
            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # настрой ширины колонок таблицы
        widths = [Mm(60), Mm(50), Mm(20), Mm(30), Mm(30), ]
        self.style_tab_set_widths_for_cols(table, widths)
        # добавь пустую строку между таблицами
        mid_paragraph = self.docx.add_paragraph(' ')
        mid_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mid_paragraph.paragraph_format.line_spacing = Pt(6)
        for i in mid_paragraph.runs:
            i.font.size = Pt(6)

    def setting_body_for_chem_tab(self):
        # взять данные для заполнения таблицы в зависимости от симметрии
        if self.si.is_symmetrical:
            head_row = self.si.chem_tab.head_row
            data_rows = self.si.chem_tab.symm_data_rows
        else:
            head_row = self.si.chem_tab.asymm_head_row[self.si.sert_number_index]
            data_rows = self.si.chem_tab.asymm_data_rows[self.si.sert_number_index]
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=len(head_row))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True)
        # заполнить тело таблицы
        for row in data_rows:
            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # настрой ширины колонок таблицы
        widths = [Mm(20), Mm(22), ]  # всего 190 Mm
        r_widths = int(148 // (len(table.columns) - 2))
        for i in range(len(table.columns) - 2):
            widths += [Mm(r_widths)]
        corr = 148 - (r_widths * (len(table.columns) - 2))
        if corr > 0:
            widths[0] = Mm(20 + corr)
        self.style_tab_set_widths_for_cols(table, widths)

    def setting_body_for_ctk_mech_tab(self):
        # взять данные для заполнения таблицы в зависимости от симметрии
        head_row = self.si.ctk_mech_tab.casts_head_row[self.si.sert_number_index]
        data_rows = self.si.ctk_mech_tab.casts_data_rows[self.si.sert_number_index]
        # создать и настроить таблицу
        table = self.docx.add_table(rows=1, cols=len(head_row))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # заполнить шапку таблицы
        hdr_cells = table.rows[0].cells
        for i in range(len(head_row)):
            hdr_cells[i].text = str(head_row[i])
        self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True,
                                                      is_change_font_size=True, font_size=11)
        # заполнить тело таблицы
        for row in data_rows:
            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False,
                                                          is_change_font_size=True, font_size=11)
        # настроить повторяющийся заголовок
        table = self.add_repeat_table_header(table)
        # настрой ширины колонок таблицы
        widths = [Mm(10), Mm(50), Mm(11), Mm(20), Mm(26),]  # всего 277 Mm
        r_widths = int(160 / (len(table.columns) - 5))
        for i in range(len(table.columns) - 5):
            widths += [Mm(r_widths)]
        corr = int(160 % (len(table.columns) - 5))
        if corr > 0:
            widths[2] = Mm(11 + corr)
        self.style_tab_set_widths_for_cols(table, widths)
        # отступ от след таблицы
        prf = self.docx.add_paragraph(' ')
        prf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prf.paragraph_format.line_spacing = Pt(6)

    def setting_body_for_ctk_chem_tab(self):
        # взять данные для заполнения таблицы в зависимости от симметрии
        head_row = self.si.ctk_mech_tab.chem_head_row[self.si.sert_number_index]
        data_rows = self.si.ctk_mech_tab.chem_data_rows[self.si.sert_number_index]
        # создать и настроить таблицу
        cols_index = len(head_row)
        outer_tab = None
        inner_tab = None
        if cols_index > 9:
            outer_tab = cols_index
        elif 9 >= cols_index > 6:
            outer_tab = 2
            inner_tab = cols_index
        elif cols_index <= 6:
            outer_tab = 3
            inner_tab = cols_index
        o_table = self.docx.add_table(rows=1, cols=outer_tab)
        o_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tab_list = []
        if outer_tab is not None:
            hdr_cells = o_table.rows[0].cells
            for i in range(outer_tab):
                self.delete_paragraph(hdr_cells[i].paragraphs[0])
                tab = hdr_cells[i].add_table(rows=1, cols=inner_tab)
                tab.style = 'Table Grid'
                tab.alignment = WD_TABLE_ALIGNMENT.LEFT
                tab_list.append(tab)
        else:
            tab_list.append(o_table)
            o_table.style = 'Table Grid'

        # заполняем шапку
        row_summ_index = len(tab_list)
        mirror_tab_iter_index = 0
        for i in range(row_summ_index):
            table = tab_list[mirror_tab_iter_index]
            # заполнить шапку таблицы
            hdr_cells = table.rows[0].cells
            for i in range(len(head_row)):
                hdr_cells[i].text = str(head_row[i])
            self.style_tabrow_set_alignment_and_font_bold(table.rows[0], font_bold=True,
                                                          is_change_font_size=True, font_size=10)
            # настроить повторяющийся заголовок
            table = self.add_repeat_table_header(table)
            mirror_tab_iter_index += 1
            # настрой ширины колонок таблицы
            widths = [Mm(16), ]  # всего 277 Mm
            for i in range(len(head_row)):
                widths += [Mm(14), ]
            self.style_tab_set_widths_for_cols(table, widths)
        # заполняем тело
        row_summ_index = len(data_rows)
        tab_iter_index = outer_tab
        mirror_tab_iter_index = 0
        for i in range(row_summ_index):
            table = tab_list[mirror_tab_iter_index]

            current_table_row = table.add_row()
            row_cells = current_table_row.cells
            for j in range(len(data_rows[i])):
                row_cells[j].text = str(data_rows[i][j])
            self.style_tabrow_set_alignment_and_font_bold(current_table_row, font_bold=False,
                                                          is_change_font_size=True, font_size=10)

            mirror_tab_iter_index += 1
            tab_iter_index -= 1
            if tab_iter_index == 0:
                tab_iter_index = outer_tab
                mirror_tab_iter_index = 0


class GroupManager(ModelDataLoader):
    def __init__(self):
        self.serts_to_print = None
        self.fatal_error = False
        self.serts_incarnations_list = []
        self.docx_list = []
        self.errors_list = []  # (level, text,)

        self.load_serts_to_print()

    def load_serts_to_print(self):
        serts = Sert.objects.filter(is_print=True)
        if not len(serts) > 0:
            self.fatal_error = True
            error_text = (f'GroupManager: ни один сертификат не выбран для печати '
                          f'(нет ни одной галочки is_print/на_печать в Sert)')
            error = (messages.ERROR, error_text)
            self.errors_list.append(error)
            return
        else:
            self.serts_to_print = serts
            self.create_serts_incarnations()
            self.create_docx_documents()
            self.fill_docx_list()

    def load_model_data(self, sert):
        NT = namedtuple('flags', ['is_attach_data', 'is_melt_data',])
        DATA_BY_TYPES = {
            'НАСОС': NT(is_attach_data=False, is_melt_data=False),
            'АРМАТУРА': NT(is_attach_data=False, is_melt_data=False),
            'НАСОС_КУСОЧКИ': NT(is_attach_data=True, is_melt_data=False),
            'АРМАТУРА_КУСОЧКИ': NT(is_attach_data=True, is_melt_data=False),
            'РЕМКОМПЛЕКТ': NT(is_attach_data=True, is_melt_data=False),
            'ГАЛЬВАНИКА': NT(is_attach_data=True, is_melt_data=False),
            'НАСОС_ХИМ': NT(is_attach_data=True, is_melt_data=True),
            'АРМАТУРА_ХИМ': NT(is_attach_data=True, is_melt_data=True),
            'СЕРТ_ЦТК': NT(is_attach_data=True, is_melt_data=True),
            'СЕРТ_ЦТК_ГГ_НАСОС': NT(is_attach_data=True, is_melt_data=True),
            'СЕРТ_ЦТК_НТ_ВЭЛВ': NT(is_attach_data=True, is_melt_data=True),  # ???
            # 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': '',
        }
        kernel_data = self.get_kernel_data(sert)
        attach_data = None
        melt_data = None
        if DATA_BY_TYPES[sert.sert_type].is_attach_data:
            attach_data = list(self.get_attachment_data(sert))
            if not attach_data:
                self.fatal_error = True
                error_text = (f'GroupManager: {sert.number_spg.number_spg}-{sert.sert_type}-{sert.number_unique.id} '
                              f'не содержит необходимых вложений в Attachment')
                error = (messages.ERROR, error_text)
                self.errors_list.append(error)
            if DATA_BY_TYPES[sert.sert_type].is_melt_data:
                melt_data = []
                for attach_item in attach_data:
                    melt = list(self.get_melt_data(attach_item))
                    if (not melt and attach_item.is_cast):
                        self.fatal_error = True
                        error_text = (
                            f'GroupManager: вложение {attach_item.designation} {attach_item.denomination} из '
                            f'{sert.number_spg.number_spg}-{sert.sert_type}-{sert.number_unique.id} '
                            f'не имеет привязанных плавок в Melt')
                        error = (messages.ERROR, error_text)
                        self.errors_list.append(error)
                    for m in melt:
                        if m not in melt_data:
                            melt_data.append(m)
        return kernel_data, attach_data, melt_data

    def create_serts_incarnations(self):
        for sert in self.serts_to_print:
            kernel_data, attach_data, melt_data = self.load_model_data(sert)
            print('========================',
                  '\nsert_type:', sert.sert_type,
                  '\nkernel_data:', kernel_data,
                  '\nattach_data:', attach_data,
                  '\nmelt_data:', melt_data,
                  '\n--------------------',)
            if not self.fatal_error:
                si = SertIncarnation(
                    sert=sert,
                    kernel_data=kernel_data,
                    attach_data=attach_data,
                    melt_data=melt_data,
                )
                # настроить style
                si = self.set_style(si)
                print("настроить style",
                    '\norganization:', si.organization,
                    '\npaging:', si.paging,
                    '\npage_orientation:', si.page_orientation,
                    '\nhead_image:', si.head_image,
                    '\nis_galv_date:', si.is_galv_date,
                    '\nmelt_anatomy:', si.melt_anatomy,
                    '\nconclusion_type:', si.conclusion_type,
                    '\nsign_type:', si.sign_type,
                    '\nhead_gost_str:', si.head_gost_str,
                    '\ntabs_names_list:', si.tabs_names_list,
                    '\n--------------------',)
                # создать tabs с данными
                si = self.set_tabs(si)
                print("создать tabs с данными",
                      '\ndocx_tab:', si.docx_tab,
                      '\nmain_tab:', si.main_tab,)
                print('main_parts_tab:')
                if si.main_parts_tab:
                    for i in si.main_parts_tab:
                        print(i)
                else:
                    print(si.main_parts_tab)
                print('parts_tab:')
                if si.parts_tab:
                    for i in si.parts_tab:
                        for ii in i:
                            if isinstance(ii, type(list())):
                                for iii in ii:
                                    print(iii)
                            else:
                                print(ii)
                else:
                    print(si.parts_tab)
                print('galv_parts_tab:')
                if si.galv_parts_tab:
                    for i in si.galv_parts_tab:
                        print(i)
                else:
                    print(si.galv_parts_tab)
                print('casts_tab:')
                if si.casts_tab:
                    for i in si.casts_tab:
                        print(i)
                else:
                    print(si.casts_tab)
                print('chem_tab:')
                if si.chem_tab:
                    for i in si.chem_tab:
                        print(i)
                else:
                    print(si.chem_tab)
                print('ctk_mech_tab:')
                if si.ctk_mech_tab:
                    for i in si.ctk_mech_tab:
                        print(i)
                else:
                    print(si.ctk_mech_tab)


                print('errors_list:')
                if si.errors_list:
                    for i in si.errors_list:
                        print(i)
                else:
                    print(si.errors_list)
                print('\n--------------------', )
                self.serts_incarnations_list.append(si)
                # куда-то надо снятие is_print, наверное в fill_docx_list

    def set_style(self, si):
        ssm = SertStyleMaker(si)
        si = ssm.get_result()
        return si

    def set_tabs(self, si):
        stm = SertTabsMaker(si)
        si = stm.get_result()
        return si

    def create_docx_documents(self):
        for si in self.serts_incarnations_list:
            docx = si.docx
            iter_index = si.iter_index
            while iter_index:
                if not docx:
                    docx = Document()
                    dm = DocxMaker(si, docx)
                    docx = dm.get_docx()
                else:
                    dm = DocxMaker(si, docx)
                    docx = dm.get_docx()
                si.sert_number_index += 1
                iter_index -= 1
            si.docx = docx
            si.sert_number_index = 1

    def fill_docx_list(self):
        docx_parts = namedtuple('docx_parts', ['content', 'name', 'format'])
        FILENAMES = {
            'НАСОС': 'PUMP',
            'АРМАТУРА': 'ARMATURE',
            'НАСОС_КУСОЧКИ': 'PUMP_PARTS',
            'АРМАТУРА_КУСОЧКИ': 'ARM_PARTS',
            'РЕМКОМПЛЕКТ': 'PARTS',
            'ГАЛЬВАНИКА': 'GALV',
            'НАСОС_ХИМ': 'PUMP_CHEM',
            'АРМАТУРА_ХИМ': 'ARM_CHEM',
            'СЕРТ_ЦТК': 'CTK_MAIN',
            'СЕРТ_ЦТК_ГГ_НАСОС': 'CTK_GG',
            'СЕРТ_ЦТК_НТ_ВЭЛВ': 'CTK_NT_VELV',
            # 'СЕРТ_ЦТК_АРМАПРОМ_20ГЛ_60': 'CTK_ARMA_PROM',
        }
        format_str = 'docx'
        for si in self.serts_incarnations_list:
            first_name_str = si.main_tab.data_row.number_spg.split('.')
            first_name_str = '-'.join(first_name_str)
            try:
                second_name_str = FILENAMES[si.sert_type]
            except KeyError:
                second_name_str = 'UNKNOWN'
            name_str = f'{first_name_str}-{second_name_str}'
            self.docx_list.append(docx_parts(
                content=si.docx,
                name=name_str,
                format=format_str,
            ))
            si.sert.is_print = False
            si.sert.save()

    def get_docx_list(self):
        if self.fatal_error is True:
            return None
        else:
            return self.docx_list

    def get_error(self):
        return self.errors_list
